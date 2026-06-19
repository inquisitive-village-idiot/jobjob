#!/usr/bin/env python3
"""Render a cover-letter body to a professional DOCX via python-docx."""

from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from jobjob.structure.applicant import Applicant

FONT_NAME = "Arial"


def create_cover_letter_docx(
    text: str,
    output_path: Path,
    role_title: str,
    company_name: str,
    applicant: Applicant,
) -> Path:
    """Write ``text`` as a formatted cover-letter DOCX.

    Arguments:
        text: The cover-letter body (paragraphs separated by blank lines).
        output_path: Destination DOCX path.
        role_title: Role title (used in metadata and the RE line).
        company_name: Company name (used in metadata).
        applicant: Applicant identity for the header.
    Returns:
        The output path.
    """
    output_path = Path(output_path)
    name = applicant.name or ""
    doc = DocxDocument()

    core_props = doc.core_properties
    core_props.author = name
    core_props.title = f"Cover Letter - {role_title}"
    core_props.subject = f"Application for {role_title} at {company_name}"
    core_props.comments = ""
    core_props.created = datetime.now()
    core_props.modified = datetime.now()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    _add_run(doc, name, size=12, bold=True, space_after=0)
    _add_run(doc, applicant.contact_line(), size=10, space_after=6)
    _add_run(doc, datetime.now().strftime("%B %d, %Y"), size=11, space_after=6)
    _add_run(doc, f"RE: {role_title}", size=11, bold=True, space_after=6)

    for para in text.strip().split("\n\n"):
        if para.strip():
            paragraph = doc.add_paragraph(para.replace("\n", " ").strip())
            paragraph.paragraph_format.space_after = Pt(12)
            # Left-aligned: full justification produces uneven word spacing at
            # this column width, which reads worse than a ragged right edge.
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = FONT_NAME

    doc.save(str(output_path))
    return output_path


def _add_run(doc, text, size, bold=False, space_after=6):
    """Append a single-run paragraph with the given size/weight/spacing."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = FONT_NAME
    paragraph.paragraph_format.space_after = Pt(space_after)
    return paragraph


# __END__
