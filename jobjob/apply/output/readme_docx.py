#!/usr/bin/env python3
"""Render the per-application README to a DOCX (uploaded to Drive as a Google Doc)."""

from collections.abc import Iterable, Mapping, Sequence
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument

from jobjob.structure.fit import Fit
from jobjob.structure.job_decription import JobDescription


def _join(values: Iterable[str]) -> str:
    return ", ".join(v for v in values if v)


def _labeled_bullet(doc, label: str, value: str) -> None:
    """Add a ``• Label: value`` bullet (bold label)."""
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value or "—")


def _add_overview(doc, job: JobDescription) -> None:
    doc.add_heading("Overview", level=1)
    rows = [
        ("Company", job.company_name),
        ("Role", job.role_title),
        ("Department", job.department),
        ("Location", _join(job.location)),
        ("Salary", job.salary),
        ("Hiring manager", job.hiring_manager),
        ("Seniority", job.seniority_level),
    ]
    for label, value in rows:
        _labeled_bullet(doc, label, value)
    if job.summary:
        doc.add_paragraph(job.summary)


def _add_resume(
    doc,
    template_name: Optional[str],
    template_archetype: Optional[str],
    resume_changes: Optional[Iterable[str]],
) -> None:
    doc.add_heading("Resume", level=1)
    if not template_name:
        doc.add_paragraph("No resume generated (Drive skipped).")
        return
    label = template_name + (f" ({template_archetype})" if template_archetype else "")
    _labeled_bullet(doc, "Template", label)

    changes = list(resume_changes or [])
    if not changes:
        _labeled_bullet(doc, "Changes", "template used as-is (no automated edits).")
        return
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run("Changes:").bold = True
    for change in changes:
        doc.add_paragraph(change, style="List Bullet 2")


def _add_two_column(
    doc, left_title: str, left: Sequence[str], right_title: str, right: Sequence[str]
) -> None:
    """Render two side-by-side bulleted columns as a bordered table."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].paragraphs[0].add_run(left_title).bold = True
    header[1].paragraphs[0].add_run(right_title).bold = True
    left = list(left) or ["—"]
    right = list(right) or ["—"]
    for i in range(max(len(left), len(right))):
        cells = table.add_row().cells
        cells[0].text = left[i] if i < len(left) else ""
        cells[1].text = right[i] if i < len(right) else ""


def _add_axis_table(doc, title: str, score: float, categories, note: str = "") -> None:
    """Render one quantitative fit axis: score line + category table."""
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{title}: ").bold = True
    paragraph.add_run(f"{score:.2f}")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, ("Category", "Score", "Notes")):
        cell.paragraphs[0].add_run(header).bold = True
    for category in categories:
        cells = table.add_row().cells
        cells[0].text = category.name
        cells[1].text = f"{category.score:.2f}"
        cells[2].text = category.note
    if note:
        doc.add_paragraph(note)


def _add_fit(doc, fit: Fit) -> None:
    doc.add_heading("Fit", level=1)
    paragraph = doc.add_paragraph()
    paragraph.add_run("Assessment: ").bold = True
    paragraph.add_run(fit.band)
    # Quantitative axes are additive: rendered only when computable, and the
    # section is byte-identical to the pre-quantitative layout when absent.
    if fit.role_fit is not None:
        _add_axis_table(
            doc, "Role fit", fit.role_fit, fit.role_fit_categories, fit.role_fit_note
        )
    elif fit.role_fit_note:
        doc.add_paragraph(fit.role_fit_note)
    if fit.preference_fit is not None:
        _add_axis_table(
            doc, "Preference fit", fit.preference_fit, fit.preference_fit_categories
        )
    _add_two_column(doc, "Strengths", fit.strengths, "Weaknesses", fit.weaknesses)


def _add_bullet_list(doc, heading: str, items: Iterable[str], level: int = 2) -> None:
    items = list(items)
    if not items:
        return
    doc.add_heading(heading, level=level)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _add_ats(doc, ats) -> None:
    """Render the ATS assessment section.

    Skipped assessments state that and nothing else; side lists (skills-file
    candidates, up-skill targets) are informational only and never feed
    generation.
    """
    doc.add_heading("ATS assessment", level=1)
    if ats.skipped:
        doc.add_paragraph("Skipped — no resume document was generated (Drive skipped).")
        return

    paragraph = doc.add_paragraph()
    paragraph.add_run("Keyword coverage: ").bold = True
    score = ats.coverage_score
    paragraph.add_run("—" if score is None else f"{score:.2f}")

    _add_bullet_list(doc, "Present in resume", ats.present)
    _add_bullet_list(doc, "Missing (evidenced)", ats.missing_evidenced)
    _add_bullet_list(doc, "Missing (unevidenced)", ats.missing_unevidenced)
    _add_bullet_list(doc, "Recommendations", ats.recommendations)
    _add_bullet_list(
        doc,
        "Skills-file candidates (evidenced but not in your skills file)",
        ats.skills_file_candidates,
    )
    _add_bullet_list(
        doc,
        "Possible up-skill targets (not for the resume or cover letter)",
        ats.upskill_targets,
    )
    _add_bullet_list(doc, "Fit vs. ATS gaps (supported but not rendered)", ats.fit_gaps)

    warnings = [f"{c.name}: {c.reason}" for c in ats.checks if not c.passed]
    doc.add_heading("Parseability", level=2)
    if warnings:
        for warning in warnings:
            doc.add_paragraph(warning, style="List Bullet")
    else:
        doc.add_paragraph("All checks passed.")


def _add_issues(doc, issues: Iterable[str]) -> None:
    doc.add_heading("Issues", level=1)
    issues = list(issues)
    if not issues:
        doc.add_paragraph("None noted.")
        return
    for issue in issues:
        doc.add_paragraph(issue, style="List Bullet")


def _add_items(
    doc, items: Iterable[Mapping], primary: str, *detail_fields: str
) -> None:
    items = list(items)
    if not items:
        doc.add_paragraph("None.", style="List Bullet")
        return
    for item in items:
        skill = item.get(primary, "")
        details = "; ".join(item.get(f, "") for f in detail_fields if item.get(f))
        text = f"{skill} — {details}" if details else skill
        doc.add_paragraph(text, style="List Bullet")


def _add_skills(doc, skills: Mapping, unmapped: Optional[Iterable[str]] = None) -> None:
    doc.add_heading("Skills analysis", level=1)

    doc.add_heading("Critical gaps", level=2)
    _add_items(
        doc, skills.get("critical_gaps", []), "skill", "why_critical", "mitigation"
    )

    unmapped = list(unmapped or [])
    if unmapped:
        doc.add_heading("Unmapped requirements", level=2)
        doc.add_paragraph(
            "Not in the skill cloud — either a genuine gap or a taxonomy hole:"
        )
        for text in unmapped:
            doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Critical (supported)", level=2)
    _add_items(doc, skills.get("critical_supported", []), "skill", "evidence")

    doc.add_heading("Important (supported)", level=2)
    _add_items(doc, skills.get("important_supported", []), "skill", "evidence")

    doc.add_heading("Strong supporting", level=2)
    _add_items(doc, skills.get("strong_supporting", []), "skill", "relevance")


def create_readme_docx(
    output_path: Path,
    job: JobDescription,
    skills: Mapping,
    fit: Fit,
    issues: Iterable[str],
    template_name: Optional[str] = None,
    template_archetype: Optional[str] = None,
    resume_changes: Optional[Iterable[str]] = None,
    unmapped: Optional[Iterable[str]] = None,
    ats: Optional[Any] = None,
) -> Path:
    """Render the application README to a DOCX.

    Arguments:
        output_path: Destination DOCX path.
        job: The parsed job description.
        skills: The skills-analysis mapping (folded in).
        fit: The fit assessment.
        issues: Processing/content issues to surface.
        template_name: Resume template used (None when Drive was skipped).
        template_archetype: Human-readable archetype of the template.
        resume_changes: Edits applied to the template (empty = used as-is).
        unmapped: JD requirements with no skill-cloud match (omitted when empty).
        ats: ATS assessment (AtsAssessment); None omits the section entirely.
    Returns:
        The output path.
    """
    output_path = Path(output_path)
    doc = DocxDocument()
    doc.add_heading(
        f"Application Summary — {job.company_name} / {job.role_title}", level=0
    )
    _add_overview(doc, job)
    _add_resume(doc, template_name, template_archetype, resume_changes)
    _add_fit(doc, fit)
    if ats is not None:
        _add_ats(doc, ats)
    _add_issues(doc, issues)
    _add_skills(doc, skills, unmapped=unmapped)
    doc.save(str(output_path))
    return output_path


# __END__
