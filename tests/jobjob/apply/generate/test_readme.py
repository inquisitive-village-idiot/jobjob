#!/usr/bin/env python3
"""Test."""

import logging
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase, mock

import jobjob.apply.generate.readme as MOD
from jobjob.structure.fit import BAND_MODERATE, BAND_STRONG, BAND_WEAK
from jobjob.structure.job_decription import JobDescription

LOGGER = logging.getLogger(__name__)


class FakeCloudSkill:
    def __init__(self, id_: str, categories: dict) -> None:
        self.id = id_
        self.categories = categories


class FakeCloud:
    """Minimal SkillCloud stand-in: exact-match resolve over a dict."""

    def __init__(self, hits: dict) -> None:
        self.hits = hits

    def resolve(self, text: str):
        return self.hits.get(text)


class ThisTestCase(TestCase):
    """Base test case for the module."""

    def patch_cloud(self, hits: dict) -> None:
        patcher = mock.patch.object(
            MOD, "get_skill_cloud", return_value=FakeCloud(hits)
        )
        patcher.start()
        self.addCleanup(patcher.stop)

    def get_tmpdir(self) -> Path:
        tmpdir = TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        return Path(tmpdir.name)

    def make_job(self, **kwargs) -> JobDescription:
        defaults = {
            f: ""
            for f in (
                "company_name",
                "role_title",
                "department",
                "seniority_level",
                "salary",
                "hiring_manager",
                "summary",
            )
        }
        defaults.update(
            {
                f: ()
                for f in (
                    "location",
                    "key_requirements",
                    "responsibilities",
                    "technical_skills",
                    "soft_skills",
                    "keywords",
                )
            }
        )
        defaults.update(kwargs)
        return JobDescription(**defaults)


class TestAssessFit(ThisTestCase):
    """Test function."""

    def test_strong_when_no_gaps(self) -> None:
        skills = {"critical_gaps": [], "critical_supported": [{"skill": "Python"}]}
        fit = MOD.assess_fit(skills)
        self.assertEqual(BAND_STRONG, fit.band)
        self.assertIn("Python", fit.summary)

    def test_moderate_with_few_gaps(self) -> None:
        skills = {"critical_gaps": [{"skill": "Rust"}, {"skill": "Go"}]}
        self.assertEqual(BAND_MODERATE, MOD.assess_fit(skills).band)

    def test_weak_with_many_gaps(self) -> None:
        skills = {"critical_gaps": [{"skill": "a"}, {"skill": "b"}, {"skill": "c"}]}
        self.assertEqual(BAND_WEAK, MOD.assess_fit(skills).band)

    def test_placeholder_scores_are_none(self) -> None:
        fit = MOD.assess_fit({"critical_gaps": []})
        self.assertIsNone(fit.role_fit)
        self.assertIsNone(fit.preference_fit)

    def test_weak_with_any_blocking_gap(self) -> None:
        skills = {
            "critical_gaps": [{"skill": "Rust", "severity": "blocking"}],
            "critical_supported": [{"skill": s} for s in "abcdefgh"],
        }
        self.assertEqual(BAND_WEAK, MOD.assess_fit(skills).band)

    def test_strong_when_only_aspirational_gaps(self) -> None:
        skills = {
            "critical_gaps": [{"skill": "Quantum", "severity": "aspirational"}],
            "critical_supported": [{"skill": "Python"}],
        }
        fit = MOD.assess_fit(skills)
        self.assertEqual(BAND_STRONG, fit.band)
        self.assertEqual(("Quantum (aspirational)",), fit.weaknesses)

    def test_moderate_when_support_outweighs_stretch_gaps(self) -> None:
        # Many stretch gaps would read Weak alone, but heavy supporting
        # evidence (>= 2x) keeps the band at Moderate.
        skills = {
            "critical_gaps": [
                {"skill": s, "severity": "stretch"} for s in ("a", "b", "c", "d")
            ],
            "critical_supported": [{"skill": s} for s in "abcdefgh"],
        }
        self.assertEqual(BAND_MODERATE, MOD.assess_fit(skills).band)

    def test_strengths_and_weaknesses_are_structured(self) -> None:
        skills = {
            "critical_gaps": [{"skill": "Rust"}],
            "critical_supported": [{"skill": "Python"}],
            "important_supported": [{"skill": "Docs"}],
        }
        fit = MOD.assess_fit(skills)
        self.assertEqual(("Python", "Docs"), fit.strengths)
        self.assertEqual(("Rust",), fit.weaknesses)


class TestScoreRoleFit(ThisTestCase):
    """Test function (deterministic role-fit scoring)."""

    HITS = {
        "Python": FakeCloudSkill("python", {"technical": 1.0}),
        "Kubernetes": FakeCloudSkill("kubernetes", {"technical": 1.0}),
        "Team leadership": FakeCloudSkill(
            "manage_a_team", {"leadership": 0.7, "collaboration": 0.3}
        ),
    }

    def test_hand_computed_category_scores(self) -> None:
        self.patch_cloud(self.HITS)
        skills = {
            "critical_supported": [{"skill": "Python"}],  # match 1.0
            "important_supported": [{"skill": "Team leadership"}],  # match 0.75
            "critical_gaps": [{"skill": "Kubernetes", "severity": "stretch"}],  # 0.4
        }

        categories, score, note = MOD._score_role_fit(skills)
        by_name = {c.name: c for c in categories}

        with self.subTest("technical = (1.0*1 + 0.4*1) / 2"):
            self.assertEqual(0.7, by_name["technical"].score)

        with self.subTest("leadership = 0.75 (sole contributor)"):
            self.assertEqual(0.75, by_name["leadership"].score)

        with self.subTest("collaboration = 0.75 (same item, fractional weight)"):
            self.assertEqual(0.75, by_name["collaboration"].score)

        with self.subTest("axis = mass-weighted mean = matched mass / total mass"):
            # weighted: technical 1.4, leadership 0.525, collaboration 0.225
            # mass:     technical 2.0, leadership 0.7,   collaboration 0.3
            expected = round((1.4 + 0.525 + 0.225) / (2.0 + 0.7 + 0.3), 2)
            self.assertEqual(expected, score)
            # a sliver-backed category must not count as much as a full one:
            unweighted_mean = round((0.7 + 0.75 + 0.75) / 3, 2)
            self.assertNotEqual(unweighted_mean, score)

        with self.subTest("full coverage -> empty note"):
            self.assertEqual("", note)

        with self.subTest("notes name contributing skills"):
            self.assertIn("Python", by_name["technical"].note)
            self.assertIn("Kubernetes", by_name["technical"].note)

    def test_sliver_category_cannot_drag_the_axis(self) -> None:
        """The motivating property of the mass-weighted axis: a weak category
        backed only by a small fractional weight barely moves the axis."""
        hits = {
            "Python": FakeCloudSkill("python", {"technical": 1.0}),
            "SQL": FakeCloudSkill("sql", {"technical": 1.0}),
            "Advising": FakeCloudSkill(
                "advising", {"communication": 0.8, "domain": 0.2}
            ),
        }
        self.patch_cloud(hits)
        skills = {
            # heavy, fully-supported technical mass...
            "critical_supported": [{"skill": "Python"}, {"skill": "SQL"}],
            # ...vs a blocking gap whose only domain presence is a 0.2 sliver
            "critical_gaps": [{"skill": "Advising", "severity": "blocking"}],
        }

        categories, score, _ = MOD._score_role_fit(skills)
        by_name = {c.name: c for c in categories}

        with self.subTest("the sliver category itself scores 0"):
            expected = 0.0
            found = by_name["domain"].score
            self.assertEqual(expected, found)

        with self.subTest("axis stays near the heavy mass, not the naive mean"):
            # mass-weighted: (2.0 + 0 + 0) / (2.0 + 0.8 + 0.2) = 0.67
            # naive mean would be (1.0 + 0 + 0) / 3 = 0.33
            expected = 0.67
            found = score
            self.assertEqual(expected, found)

    def test_reproducible(self) -> None:
        self.patch_cloud(self.HITS)
        skills = {
            "critical_supported": [{"skill": "Python"}],
            "critical_gaps": [{"skill": "Kubernetes", "severity": "blocking"}],
        }
        expected = MOD._score_role_fit(skills)
        found = MOD._score_role_fit(skills)
        self.assertEqual(expected, found)

    def test_aspirational_and_noncanonical_excluded(self) -> None:
        self.patch_cloud(self.HITS)
        skills = {
            "critical_supported": [{"skill": "Python"}],
            "critical_gaps": [
                {"skill": "Quantum", "severity": "aspirational"},  # never scored
                {"skill": "Alpaca wrangling", "severity": "stretch"},  # no cloud hit
            ],
        }

        categories, score, note = MOD._score_role_fit(skills)

        with self.subTest("only the canonical supported skill scores"):
            self.assertEqual(1.0, score)

        with self.subTest("coverage note counts the exclusion"):
            self.assertIn("1 of 2", note)
            self.assertIn("Alpaca wrangling", note)

        with self.subTest("aspirational is not even counted as excluded"):
            self.assertNotIn("Quantum", note)

    def test_zero_coverage_returns_none_with_note(self) -> None:
        self.patch_cloud({})
        skills = {"critical_supported": [{"skill": "Anything"}]}

        categories, score, note = MOD._score_role_fit(skills)

        self.assertEqual((), categories)
        self.assertIsNone(score)
        self.assertIn("1 of 1", note)

    def test_no_items_returns_none_silently(self) -> None:
        self.patch_cloud(self.HITS)
        expected = ((), None, "")
        found = MOD._score_role_fit({})
        self.assertEqual(expected, found)

    def test_cloud_unavailable_degrades(self) -> None:
        patcher = mock.patch.object(
            MOD, "get_skill_cloud", side_effect=FileNotFoundError("no cloud")
        )
        patcher.start()
        self.addCleanup(patcher.stop)

        skills = {"critical_supported": [{"skill": "Python"}]}
        expected = ((), None, "")
        found = MOD._score_role_fit(skills)
        self.assertEqual(expected, found)

    def test_band_unchanged_by_scoring(self) -> None:
        self.patch_cloud(self.HITS)
        skills = {
            "critical_gaps": [{"skill": "Kubernetes", "severity": "blocking"}],
            "critical_supported": [{"skill": "Python"}],
        }
        fit = MOD.assess_fit(skills)
        self.assertEqual(BAND_WEAK, fit.band)  # blocking gap rule, as before


class TestParsePreferenceFit(ThisTestCase):
    """Test function (model-judged preference-fit parsing)."""

    def test_parses_valid_block(self) -> None:
        skills = {
            "preference_fit": [
                {"name": "role type", "score": 0.8, "note": "matches goals"},
                {"name": "location", "score": 0.4, "note": "relocation needed"},
            ]
        }

        categories, score = MOD._parse_preference_fit(skills)

        with self.subTest("categories parsed"):
            found = tuple(c.name for c in categories)
            self.assertEqual(("role type", "location"), found)

        with self.subTest("axis is the mean"):
            self.assertEqual(0.6, score)

    def test_missing_block_returns_none(self) -> None:
        expected = ((), None)
        found = MOD._parse_preference_fit({})
        self.assertEqual(expected, found)

    def test_malformed_categories_skipped_with_warning(self) -> None:
        skills = {
            "preference_fit": [
                {"name": "ok", "score": 0.5},
                {"name": "no score"},
                {"name": "out of range", "score": 1.5},
                {"name": "non-numeric", "score": "high"},
                "not even a mapping",
            ]
        }

        with self.assertLogs(
            "jobjob.apply.generate.readme", level="WARNING"
        ) as captured:
            categories, score = MOD._parse_preference_fit(skills)

        with self.subTest("only the valid category survives"):
            self.assertEqual(1, len(categories))
            self.assertEqual(0.5, score)

        with self.subTest("each invalid category warned"):
            self.assertEqual(4, len(captured.records))

    def test_all_invalid_returns_none(self) -> None:
        skills = {"preference_fit": [{"name": "bad", "score": 2.0}]}
        with self.assertLogs("jobjob.apply.generate.readme", level="WARNING"):
            expected = ((), None)
            found = MOD._parse_preference_fit(skills)
        self.assertEqual(expected, found)


class TestGenerateApplicationReadme(ThisTestCase):
    """Test function."""

    def test_writes_nonempty_docx(self) -> None:
        job = self.make_job(company_name="Acme", role_title="Engineer")
        skills = {
            "critical_gaps": [
                {"skill": "Rust", "why_critical": "x", "mitigation": "y"}
            ],
            "critical_supported": [{"skill": "Python", "evidence": "resume"}],
            "important_supported": [],
            "strong_supporting": [],
        }
        out = Path(self.get_tmpdir(), "readme.docx")
        result = MOD.generate_application_readme(
            job,
            skills,
            out,
            issues=["parse warning"],
            template_name="features_writer",
            template_archetype="Features Writer",
            resume_changes=["Swapped objective for the role"],
        )
        self.assertEqual(out, result)
        self.assertTrue(out.is_file())
        self.assertGreater(out.stat().st_size, 0)

    def test_renders_template_and_two_column_fit(self) -> None:
        """README shows the template, its changes, and a Strengths/Weaknesses table."""
        from docx import Document as DocxDocument

        job = self.make_job(company_name="Acme", role_title="Engineer")
        skills = {
            "critical_gaps": [{"skill": "Rust"}],
            "critical_supported": [{"skill": "Python"}],
        }
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(
            job,
            skills,
            out,
            template_name="features_writer",
            template_archetype="Features Writer",
            resume_changes=["Objective swap"],
        )
        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("features_writer", text)
        self.assertIn("Features Writer", text)
        self.assertIn("Objective swap", text)
        # Strengths/Weaknesses rendered as a two-column table (located by
        # header: quantitative axis tables may render before it).
        self.assertTrue(doc.tables, "expected a fit table")
        headers = [[c.text for c in t.rows[0].cells] for t in doc.tables]
        self.assertIn(["Strengths", "Weaknesses"], headers)
        sw_table = doc.tables[headers.index(["Strengths", "Weaknesses"])]
        table_text = "\n".join(c.text for r in sw_table.rows for c in r.cells)
        self.assertIn("Python", table_text)
        self.assertIn("Rust", table_text)

    def test_renders_two_axis_fit_tables(self) -> None:
        from docx import Document as DocxDocument

        self.patch_cloud({"Python": FakeCloudSkill("python", {"technical": 1.0})})
        job = self.make_job(company_name="Acme", role_title="Engineer")
        skills = {
            "critical_supported": [{"skill": "Python"}],
            "preference_fit": [
                {"name": "role type", "score": 0.8, "note": "matches goals"}
            ],
        }
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(job, skills, out)

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)

        with self.subTest("axis score lines present"):
            self.assertIn("Role fit", text)
            self.assertIn("Preference fit", text)

        with self.subTest("category tables rendered"):
            all_cells = "\n".join(
                c.text for t in doc.tables for r in t.rows for c in r.cells
            )
            self.assertIn("technical", all_cells)
            self.assertIn("role type", all_cells)
            self.assertIn("matches goals", all_cells)

    def test_readme_unchanged_when_axes_absent(self) -> None:
        """No fit block + no canonical coverage -> today's layout exactly."""
        from docx import Document as DocxDocument

        self.patch_cloud({})
        job = self.make_job(company_name="Acme", role_title="Engineer")
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(job, {"critical_gaps": []}, out)

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("Role fit", text)
        self.assertNotIn("Preference fit", text)
        # the strengths/weaknesses table is the only table
        self.assertEqual(1, len(doc.tables))

    def test_renders_ats_section(self) -> None:
        from docx import Document as DocxDocument

        from jobjob.apply.generate.ats import AtsAssessment, AtsCheck

        self.patch_cloud({})
        ats = AtsAssessment(
            coverage_score=0.64,
            present=("Python",),
            missing_evidenced=("SQL",),
            recommendations=("SQL is supported by your documentation but absent.",),
            skills_file_candidates=("Docker",),
            upskill_targets=("Kubernetes",),
            checks=(
                AtsCheck(name="content-in-tables", passed=False, reason="1 table"),
            ),
            fit_gaps=("SQL",),
        )
        job = self.make_job(company_name="Acme", role_title="Engineer")
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(job, {}, out, ats=ats)

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        for expected in (
            "ATS assessment",
            "0.64",
            "SQL is supported",
            "Skills-file candidates",
            "up-skill targets",
            "content-in-tables: 1 table",
        ):
            with self.subTest(expected=expected):
                self.assertIn(expected, text)

    def test_renders_ats_skipped(self) -> None:
        from docx import Document as DocxDocument

        from jobjob.apply.generate.ats import AtsAssessment

        self.patch_cloud({})
        job = self.make_job(company_name="Acme", role_title="Engineer")
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(job, {}, out, ats=AtsAssessment(skipped=True))

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Skipped — no resume document", text)
        self.assertNotIn("Keyword coverage", text)

    def test_omits_ats_section_when_none(self) -> None:
        from docx import Document as DocxDocument

        self.patch_cloud({})
        job = self.make_job(company_name="Acme", role_title="Engineer")
        out = Path(self.get_tmpdir(), "readme.docx")
        MOD.generate_application_readme(job, {}, out)

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("ATS assessment", text)

    def test_renders_unmapped_requirements(self) -> None:
        """JD requirements with no skill-cloud match surface in the README."""
        from unittest import mock

        from docx import Document as DocxDocument

        from jobjob.structure.normalize import NormalizedRequirement

        job = self.make_job(
            company_name="Acme",
            role_title="Engineer",
            key_requirements=("alpaca wrangling",),
        )
        normalized = (NormalizedRequirement(text="alpaca wrangling"),)
        out = Path(self.get_tmpdir(), "readme.docx")
        with mock.patch.object(MOD, "normalize_requirements", return_value=normalized):
            MOD.generate_application_readme(job, {}, out)

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Unmapped requirements", text)
        self.assertIn("alpaca wrangling", text)

    def test_omits_unmapped_section_when_cloud_unavailable(self) -> None:
        from unittest import mock

        from docx import Document as DocxDocument

        job = self.make_job(
            company_name="Acme",
            role_title="Engineer",
            key_requirements=("anything",),
        )
        out = Path(self.get_tmpdir(), "readme.docx")
        with mock.patch.object(
            MOD, "normalize_requirements", side_effect=FileNotFoundError("no cloud")
        ):
            MOD.generate_application_readme(job, {}, out)  # should not raise

        doc = DocxDocument(str(out))
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertNotIn("Unmapped requirements", text)


# __END__
