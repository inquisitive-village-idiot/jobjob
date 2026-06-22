#!/usr/bin/env python3
"""Test."""

import logging
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest import TestCase

import jobjob.loader.loadstatic as MOD
from tests.fixtures import fixture_path

LOGGER = logging.getLogger(__name__)


class ThisTestCase(TestCase):
    """Base test case for the module."""

    def get_tmpdir(self) -> Path:
        tmpdir = TemporaryDirectory()
        self.addCleanup(tmpdir.cleanup)
        return Path(tmpdir.name)


class TestLoadPdfContent(ThisTestCase):
    """Test function."""

    def test_raises_if_unable_to_extract_text(self) -> None:
        tests = [
            fixture_path("job_description_blank"),
            fixture_path("job_description_image"),
        ]
        for given in tests:
            with self.subTest(given):
                with self.assertRaisesRegex(ValueError, "not extract text"):
                    MOD.load_pdf_text(given)

    def test_returns_expected(self) -> None:
        given = fixture_path("job_description_acme")
        expected = "The Acme Gazette"
        result = MOD.load_pdf_text(given)
        found = result[: len(expected)]
        self.assertEqual(expected, found)


class TestLoadText(ThisTestCase):
    """Test function."""

    def test_reads_plain_text(self) -> None:
        given = Path(self.get_tmpdir(), "note.md")
        given.write_text("# Heading\n\nbody")
        self.assertEqual("# Heading\n\nbody", MOD.load_text(given))


class TestLoadPdfTextOrNone(ThisTestCase):
    """Test function."""

    def test_returns_text_for_text_pdf(self) -> None:
        given = fixture_path("job_description_acme")
        result = MOD.load_pdf_text_or_none(given)
        self.assertIsNotNone(result)
        self.assertIsInstance(result, str)

    def test_returns_none_for_image_only_pdf(self) -> None:
        given = fixture_path("job_description_image")
        result = MOD.load_pdf_text_or_none(given)
        self.assertIsNone(result)

    def test_returns_none_for_blank_pdf(self) -> None:
        given = fixture_path("job_description_blank")
        result = MOD.load_pdf_text_or_none(given)
        self.assertIsNone(result)


class TestLoadDocxText(ThisTestCase):
    """Test function."""

    def test_extracts_paragraphs_from_docx(self) -> None:
        from docx import Document as DocxDocument

        docx_path = Path(self.get_tmpdir(), "test.docx")
        doc = DocxDocument()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        result = MOD.load_docx_text(docx_path)

        self.assertIn("First paragraph", result)
        self.assertIn("Second paragraph", result)

    def test_skips_empty_paragraphs(self) -> None:
        from docx import Document as DocxDocument

        docx_path = Path(self.get_tmpdir(), "test.docx")
        doc = DocxDocument()
        doc.add_paragraph("")
        doc.add_paragraph("Real content")
        doc.add_paragraph("")
        doc.save(str(docx_path))

        result = MOD.load_docx_text(docx_path)

        self.assertIn("Real content", result)
        # Empty paragraphs should not add blank separators
        self.assertNotIn("\n\n\n", result)


class TestReadDocument(ThisTestCase):
    """Test dispatcher."""

    def test_dispatches_text(self) -> None:
        given = Path(self.get_tmpdir(), "a.txt")
        given.write_text("hello")
        self.assertEqual("hello", MOD.read_document(given))

    def test_dispatches_markdown(self) -> None:
        given = Path(self.get_tmpdir(), "notes.md")
        given.write_text("# heading")
        self.assertEqual("# heading", MOD.read_document(given))

    def test_dispatches_pdf(self) -> None:
        given = fixture_path("job_description_acme")
        result = MOD.read_document(given)
        self.assertTrue(result)

    def test_dispatches_docx(self) -> None:
        from docx import Document as DocxDocument

        docx_path = Path(self.get_tmpdir(), "a.docx")
        doc = DocxDocument()
        doc.add_paragraph("docx content")
        doc.save(str(docx_path))

        result = MOD.read_document(docx_path)
        self.assertIn("docx content", result)

    def test_returns_empty_for_unsupported(self) -> None:
        given = Path(self.get_tmpdir(), "a.rtf")
        given.write_text("hello")
        self.assertEqual("", MOD.read_document(given))

    def test_case_insensitive_suffix(self) -> None:
        given = Path(self.get_tmpdir(), "a.TXT")
        given.write_text("upper suffix")
        self.assertEqual("upper suffix", MOD.read_document(given))


# __END__
